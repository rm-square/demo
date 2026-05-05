import io
import os
import re
import glob
import json
import shutil
import threading
import subprocess
import anthropic
from flask import Flask, render_template, request, Response, stream_with_context, jsonify

app = Flask(__name__)

HARTMANN_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'hartmann')

SYSTEM_PROMPT = """Du bist der Revenue Intelligence Assistant von Hartmann Antriebstechnik GmbH.

Du hast Zugriff auf alle internen Dokumente: Produktinformationen, Preislisten, Vertragsklauseln, Delegation of Authority, Sales Plays, Battle Cards, Kundenaccounts, CS-Eskalationen, Marketing-Daten und Pipeline.

Du beantwortest nicht nur Fragen — du lieferst konkrete Execution:
- Dokumente erstellen: RFP-Antworten, Briefings, Analysen
- Entscheidungen vorbereiten: Pricing, Freigaben, Verhandlungsposition
- Konflikte aufzeigen: wenn eine Information aus Dokument A mit einer Anforderung aus Dokument B kollidiert
- Externe Dokumente analysieren: wenn der Nutzer ein "EXTERNES DOKUMENT" mitschickt, dieses gegen die interne Wissensbasis halten und Konflikte/Risiken identifizieren

Regeln:
- Immer konkrete Zahlen und Fakten aus den Dokumenten nennen
- Wenn eine Freigabe laut DoA nötig ist, das explizit kennzeichnen mit ⚠️
- Konflikte zwischen Dokumenten fett markieren mit 🚨
- Strukturiert antworten: Überschriften, Tabellen, Aufzählungen — kein reiner Fließtext
- Auf Deutsch antworten

QUELLEN-FORMAT (verbindlich, immer als zweitletzte Zeile der Antwort):
**Quellen:** pfad/dokument-1, pfad/dokument-2, pfad/dokument-3

- Nutze die exakten Pfade wie sie in den `=== DOKUMENT: ... ===`-Headern stehen, ohne `.md`-Endung
- Komma-separiert in EINER Zeile, kein Bullet, keine zusätzliche Formatierung
- Nur tatsächlich für die Antwort verwendete Dokumente listen

KONFLIKT-FORMAT (nur wenn direkte Widersprüche zwischen zwei internen Dokumenten erkannt wurden):
**Konflikte:** pfad/dokument-A <-> pfad/dokument-B, pfad/dokument-C <-> pfad/dokument-D

- Immer direkt nach **Quellen:**
- Nur dokumentierte Konflikte: ein Fakt aus Dokument A widerspricht direkt einem Fakt aus Dokument B
- Gleiche Pfad-Konvention wie **Quellen:** (ohne .md-Endung)
- Nur weglassen wenn keine direkten Dokumentenkonflikte vorliegen

BRAIN-SPEICHERN-FORMAT (NUR wenn der Nutzer explizit sagt "speichere das", "ins Brain schreiben", "zurückschreiben" o.ä.):
Schreibe nach der Antwort folgenden Block — er wird nicht angezeigt, sondern direkt als Dokument gespeichert:

---BRAIN-DOKUMENT-START---
pfad: kategorie/dokumentname-datum
titel: Vollständiger Dokumenttitel
---
[Vollständiger Markdown-Inhalt des Dokuments — strukturiert, mit Datum, Typ, allen relevanten Infos]
---BRAIN-DOKUMENT-ENDE---

- Kategorie muss einem bestehenden Unterordner entsprechen: commercial, company, contracts, customers, marketing, products, sales
- Dokumentname: lowercase, Bindestriche statt Leerzeichen, kein .md
- Der Inhalt soll als eigenständiges Dokument nutzbar sein — vollständig, strukturiert, mit Frontmatter-artigen Metadaten oben"""


# ─── Dynamic knowledge base — reloads after new docs are saved ──────────────
_kb_lock = threading.Lock()
_kb_cache: dict = {'context': None, 'sources': None}


def load_knowledge_base():
    docs = {}
    pattern = os.path.join(HARTMANN_DIR, '**', '*.md')
    for path in sorted(glob.glob(pattern, recursive=True)):
        filename = os.path.basename(path)
        if filename.startswith('_'):
            continue
        rel = os.path.relpath(path, HARTMANN_DIR)
        with open(path, encoding='utf-8') as f:
            docs[rel] = f.read()
    return docs


def build_context(docs):
    parts = []
    for name, content in docs.items():
        parts.append(f"=== DOKUMENT: {name} ===\n{content}")
    return "\n\n".join(parts)


def get_kb():
    with _kb_lock:
        if _kb_cache['context'] is None:
            docs = load_knowledge_base()
            _kb_cache['context'] = build_context(docs)
            _kb_cache['sources'] = {os.path.splitext(p)[0] for p in docs.keys()}
        return _kb_cache['context'], _kb_cache['sources']


def invalidate_kb():
    with _kb_lock:
        _kb_cache['context'] = None
        _kb_cache['sources'] = None


# ─── Parsing helpers ─────────────────────────────────────────────────────────
SOURCES_LINE   = re.compile(r'\*\*Quellen:\*\*\s*(.+?)\s*$', re.MULTILINE)
CONFLICTS_LINE = re.compile(r'\*\*Konflikte:\*\*\s*(.+?)\s*$', re.MULTILINE)
SAVE_BLOCK_RE  = re.compile(
    r'---BRAIN-DOKUMENT-START---\s*\n'
    r'pfad:\s*(.+?)\s*\n'
    r'titel:\s*(.+?)\s*\n'
    r'---\s*\n'
    r'([\s\S]*?)'
    r'---BRAIN-DOKUMENT-ENDE---',
    re.DOTALL
)


def extract_sources(text: str, known: set) -> list[str]:
    matches = SOURCES_LINE.findall(text)
    if not matches:
        return []
    raw = matches[-1]
    candidates = [c.strip().strip('`').rstrip('.,;').replace('.md', '')
                  for c in raw.split(',')]
    return [c for c in candidates if c in known]


def extract_conflicts(text: str, known: set) -> list[list[str]]:
    matches = CONFLICTS_LINE.findall(text)
    if not matches:
        return []
    raw = matches[-1]
    pairs = []
    for part in raw.split(','):
        if '<->' in part:
            sides = [s.strip().strip('`').rstrip('.,;').replace('.md', '')
                     for s in part.split('<->')]
            if len(sides) == 2 and sides[0] in known and sides[1] in known:
                pairs.append(sides)
    return pairs


def extract_save_block(text: str):
    m = SAVE_BLOCK_RE.search(text)
    if not m:
        return None
    return {'path': m.group(1).strip(), 'titel': m.group(2).strip(),
            'content': m.group(3).strip()}


# ─── Engine: Anthropic API ────────────────────────────────────────────────────
def stream_via_api(message, history=None):
    kb, _ = get_kb()
    client = anthropic.Anthropic(api_key=os.environ['ANTHROPIC_API_KEY'])
    messages = [{"role": h["role"], "content": h["content"]} for h in (history or [])]
    messages.append({"role": "user", "content": message})
    full_text = ""
    with client.messages.stream(
        model="claude-opus-4-7",
        max_tokens=6000,
        thinking={"type": "adaptive"},
        system=[
            {"type": "text", "text": SYSTEM_PROMPT},
            {"type": "text",
             "text": f"WISSENSBASIS HARTMANN ANTRIEBSTECHNIK:\n\n{kb}",
             "cache_control": {"type": "ephemeral"}}
        ],
        messages=messages
    ) as stream:
        for event in stream:
            if event.type == "thinking":
                yield ("thinking", event.thinking)
            elif event.type == "text":
                full_text += event.text
                yield ("text", event.text)
    yield ("full", full_text)


# ─── Engine: claude CLI subprocess ────────────────────────────────────────────
def stream_via_cli(message, history=None):
    kb, _ = get_kb()
    full_system = f"{SYSTEM_PROMPT}\n\nWISSENSBASIS HARTMANN ANTRIEBSTECHNIK:\n\n{kb}"
    if history:
        parts = [
            f"[{'Nutzer' if h['role'] == 'user' else 'Assistent'}]:\n{h['content']}"
            for h in history
        ]
        message = (f"GESPRÄCHSVERLAUF:\n\n" + "\n\n".join(parts) +
                   f"\n\n---\n\nAKTUELLE ANFRAGE:\n{message}")

    cli = shutil.which("claude")
    if not cli:
        raise RuntimeError(
            "claude CLI not found in PATH. Install Claude Code and run `claude /login`, "
            "or set ANTHROPIC_API_KEY to use the API engine."
        )

    cmd = [
        cli, "-p",
        "--verbose",
        "--model", "claude-opus-4-7",
        "--effort", "high",
        "--output-format", "stream-json",
        "--include-partial-messages",
        "--no-session-persistence",
        "--disable-slash-commands",
        "--system-prompt", full_system,
        message,
    ]

    env = {k: v for k, v in os.environ.items() if k != "ANTHROPIC_API_KEY"}
    proc = subprocess.Popen(
        cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
        text=True, bufsize=1, env=env,
    )

    full_text = ""
    try:
        for line in proc.stdout:
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
            except json.JSONDecodeError:
                continue

            if obj.get("type") == "stream_event":
                evt = obj.get("event", {})
                if evt.get("type") == "content_block_delta":
                    delta = evt.get("delta", {})
                    dtype = delta.get("type")
                    if dtype == "thinking_delta":
                        chunk = delta.get("thinking", "")
                        if chunk:
                            yield ("thinking", chunk)
                    elif dtype == "text_delta":
                        chunk = delta.get("text", "")
                        if chunk:
                            full_text += chunk
                            yield ("text", chunk)
            elif obj.get("type") == "result" and obj.get("is_error"):
                err = obj.get("result") or "claude CLI returned an error"
                raise RuntimeError(err)
    finally:
        proc.wait(timeout=30)
        if proc.returncode not in (0, None):
            stderr = proc.stderr.read() if proc.stderr else ""
            if not full_text:
                raise RuntimeError(f"claude CLI exited {proc.returncode}: {stderr.strip()[:400]}")

    yield ("full", full_text)


def select_engine():
    if os.environ.get('ANTHROPIC_API_KEY'):
        return stream_via_api, "api"
    return stream_via_cli, "cli"


# ─── Routes ───────────────────────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/graph')
def graph():
    return render_template('graph.html')


@app.route('/api/extract-doc', methods=['POST'])
def extract_doc():
    f = request.files.get('file')
    if not f:
        return jsonify({'error': 'Keine Datei empfangen'}), 400
    filename = f.filename or ''
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext in ('.txt', '.md', '.csv'):
            text = f.read().decode('utf-8', errors='replace')
        elif ext == '.docx':
            import docx as _docx
            doc = _docx.Document(io.BytesIO(f.read()))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            text = '\n'.join(parts)
        elif ext == '.pdf':
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(f.read()))
            text = '\n'.join(page.extract_text() or '' for page in reader.pages)
        else:
            return jsonify({'error': f'Format nicht unterstützt: {ext}'}), 400
        return jsonify({'text': text, 'filename': filename, 'chars': len(text)})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/save-doc', methods=['POST'])
def save_doc():
    body = request.get_json()
    path    = (body.get('path') or '').strip()
    content = (body.get('content') or '').strip()
    if not path or not content:
        return jsonify({'error': 'Pfad und Inhalt erforderlich'}), 400
    if '..' in path or path.startswith('/') or path.startswith('\\'):
        return jsonify({'error': 'Ungültiger Pfad'}), 400

    full_path = os.path.join(HARTMANN_DIR, path + '.md')
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, 'w', encoding='utf-8') as fh:
        fh.write(content)

    invalidate_kb()
    _, new_sources = get_kb()  # rebuild now so next request is fast

    return jsonify({'ok': True, 'path': path,
                    'doc_count': len(new_sources),
                    'display_path': os.path.relpath(full_path, HARTMANN_DIR)})


@app.route('/api/chat', methods=['POST'])
def chat():
    body = request.get_json()
    message      = body.get('message', '').strip()
    attached_doc = body.get('attached_doc', '').strip()
    history      = body.get('history', [])
    if not message:
        return Response('data: {"type":"error","content":"Leere Anfrage"}\n\n',
                        mimetype='text/event-stream')

    if attached_doc:
        full_message = (f"EXTERNES DOKUMENT (vom Kunden / von außen eingereicht):\n\n"
                        f"{attached_doc}\n\n---\n\nANFRAGE: {message}")
    else:
        full_message = message

    engine_fn, engine_name = select_engine()

    def generate():
        try:
            yield f"data: {json.dumps({'type': 'engine', 'content': engine_name})}\n\n"
            full_text = ""
            for kind, payload in engine_fn(full_message, history=history):
                if kind == "full":
                    full_text = payload
                else:
                    yield f"data: {json.dumps({'type': kind, 'content': payload})}\n\n"

            _, known = get_kb()
            yield f"data: {json.dumps({'type': 'sources',   'content': extract_sources(full_text, known)})}\n\n"
            conflicts = extract_conflicts(full_text, known)
            if conflicts:
                yield f"data: {json.dumps({'type': 'conflicts', 'content': conflicts})}\n\n"
            save_data = extract_save_block(full_text)
            if save_data:
                yield f"data: {json.dumps({'type': 'saveable',  'content': save_data})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'}
    )


if __name__ == '__main__':
    _, engine_name = select_engine()
    print(f"[demo] auth engine: {engine_name}  "
          f"({'API key' if engine_name == 'api' else 'claude CLI / Subscription'})")
    _, sources = get_kb()
    print(f"[demo] knowledge base: {len(sources)} Dokumente geladen")
    app.run(debug=False, port=5001)
